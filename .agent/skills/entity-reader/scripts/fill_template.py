#!/usr/bin/env python3
"""
fill_template.py
----------------
Generic template filler for entity metadata - works with any Excel or Word template format.

This script intelligently detects template structure and maps entity fields to template
columns/sections automatically. Supports both Excel (.xlsx) and Word (.docx) templates.

Usage:
    from scripts.fill_template import fill_template
    
    result = fill_template(
        entities=entities,
        template_path='path/to/template.xlsx',
        output_path='path/to/output.xlsx',
        user_column_mapping=None  # Optional: specify if auto-detection fails
    )
"""

from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime
from typing import List, Dict, Optional, Tuple
import re
from difflib import SequenceMatcher


class TemplateMapper:
    """Intelligent mapper from entity fields to template columns."""
    
    # Fuzzy matching keywords for common column types
    COLUMN_PATTERNS = {
        'field_name': ['field', 'attribute', 'property', 'name', 'column name', 'field name'],
        'column_name': ['column', 'col name', 'db column', 'database column', 'db name'],
        'sql_type': ['type', 'data type', 'datatype', 'sql type', 'db type'],
        'pk': ['pk', 'primary', 'key', 'primary key', 'is pk', 'is primary'],
        'nullable': ['null', 'nullable', 'allow null', 'not null', 'nullability', 'is null'],
        'unique': ['unique', 'distinct', 'is unique', 'unique key'],
        'relation': ['fk', 'foreign', 'relation', 'table link', 'reference', 'foreign key', 'ref table'],
        'length': ['length', 'size', 'max length', 'maxlength'],
        'default': ['default', 'data default', 'default value', 'initial value'],
        'description': ['description', 'note', 'comment', 'remark', 'notes', 'desc'],
        'validation': ['validation', 'constraint', 'rule', 'check'],
    }
    
    @staticmethod
    def similarity(a: str, b: str) -> float:
        """Calculate similarity between two strings."""
        return SequenceMatcher(None, a.lower().strip(), b.lower().strip()).ratio()
    
    @classmethod
    def fuzzy_match_column(cls, header: str, threshold: float = 0.6) -> Optional[str]:
        """
        Match a template column header to a known field type using fuzzy matching.
        
        Args:
            header: Column header from template
            threshold: Minimum similarity score (0.0 to 1.0)
        
        Returns:
            Matched field type or None
        """
        if not header or not isinstance(header, str):
            return None
        
        best_match = None
        best_score = threshold
        
        for field_type, patterns in cls.COLUMN_PATTERNS.items():
            for pattern in patterns:
                score = cls.similarity(header, pattern)
                if score > best_score:
                    best_score = score
                    best_match = field_type
        
        return best_match
    
    @classmethod
    def detect_excel_structure(cls, ws) -> Dict:
        """
        Auto-detect Excel template structure.
        
        Returns dict with:
            - header_row: Row index of headers
            - column_mapping: Dict mapping column indices to field types
            - data_start_row: Where data rows begin
            - structure_type: 'single_sheet' or 'multi_sheet'
        """
        structure = {
            'header_row': None,
            'column_mapping': {},
            'data_start_row': None,
            'structure_type': 'single_sheet'
        }
        
        # Scan first 10 rows to find headers
        for row_idx in range(1, min(11, ws.max_row + 1)):
            row_cells = [ws.cell(row=row_idx, column=col).value 
                        for col in range(1, ws.max_column + 1)]
            
            # A header row typically has multiple non-empty text values
            non_empty = [c for c in row_cells if c and isinstance(c, str)]
            if len(non_empty) >= 3:  # At least 3 columns
                # Try to match these to our known patterns
                matches = sum(1 for cell in non_empty if cls.fuzzy_match_column(cell))
                if matches >= 2:  # At least 2 recognizable columns
                    structure['header_row'] = row_idx
                    structure['data_start_row'] = row_idx + 1
                    
                    # Build column mapping
                    for col_idx, cell_value in enumerate(row_cells, start=1):
                        field_type = cls.fuzzy_match_column(cell_value)
                        if field_type:
                            structure['column_mapping'][col_idx] = field_type
                    
                    break
        
        return structure
    
    @classmethod
    def detect_word_structure(cls, doc: Document) -> Dict:
        """
        Auto-detect Word template structure.
        
        Returns dict with:
            - format: 'table' or 'sections'
            - table_index: Index of the table containing field data (if format='table')
            - column_mapping: Dict mapping column indices to field types
        """
        structure = {
            'format': None,
            'table_index': None,
            'column_mapping': {}
        }
        
        # Check for tables
        for idx, table in enumerate(doc.tables):
            if table.rows:
                # Check first row for headers
                header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                matches = sum(1 for cell in header_cells if cls.fuzzy_match_column(cell))
                
                if matches >= 2:
                    structure['format'] = 'table'
                    structure['table_index'] = idx
                    
                    # Build column mapping
                    for col_idx, cell_text in enumerate(header_cells):
                        field_type = cls.fuzzy_match_column(cell_text)
                        if field_type:
                            structure['column_mapping'][col_idx] = field_type
                    
                    break
        
        # If no table found, assume sections format
        if not structure['format']:
            structure['format'] = 'sections'
        
        return structure


def fill_excel_template(entities: List[Dict], template_path: str, output_path: str,
                       user_column_mapping: Optional[Dict] = None) -> str:
    """
    Fill an Excel template with entity metadata.
    
    Args:
        entities: List of entity dicts with 'name', 'table_name', 'fields'
        template_path: Path to template .xlsx file
        output_path: Path to save filled template
        user_column_mapping: Optional manual column mapping if auto-detection fails
    
    Returns:
        Path to saved file
    """
    
    # Load template
    wb = load_workbook(template_path, data_only=False)
    
    # Get first sheet as the working sheet (or create one if empty workbook)
    if wb.sheetnames:
        ws = wb.active
    else:
        ws = wb.create_sheet("Entities")
    
    # Detect structure
    structure = TemplateMapper.detect_excel_structure(ws)
    
    # Use user mapping if provided
    if user_column_mapping:
        structure['column_mapping'] = user_column_mapping
    
    if not structure['header_row']:
        raise ValueError(
            "Could not auto-detect template structure. Please provide user_column_mapping. "
            "Example: {'field_name': 1, 'sql_type': 2, 'nullable': 3}"
        )
    
    # Determine if we're doing single-sheet (all entities) or multi-sheet (one per entity)
    if len(entities) > 3:
        # Multi-sheet approach: one sheet per entity
        for entity in entities:
            sheet_name = entity.get('table_name', entity['name'])[:31]  # Excel max sheet name length
            
            # Create or get sheet
            if sheet_name in wb.sheetnames:
                entity_ws = wb[sheet_name]
                # Find last row
                last_row = entity_ws.max_row
                start_row = last_row + 1
            else:
                entity_ws = wb.create_sheet(title=sheet_name)
                # Copy header structure from template
                if structure['header_row']:
                    for col_idx, field_type in structure['column_mapping'].items():
                        # Copy header
                        header_cell = ws.cell(row=structure['header_row'], column=col_idx)
                        entity_ws.cell(row=1, column=col_idx, value=header_cell.value)
                start_row = 2
            
            # Fill data
            _fill_entity_rows(entity_ws, entity, structure['column_mapping'], start_row)
    
    else:
        # Single-sheet approach: all entities in one sheet
        current_row = structure['data_start_row']
        
        for entity in entities:
            # Optionally add entity name row
            if 'field_name' in structure['column_mapping'].values():
                ws.cell(row=current_row, column=1, value=f"Entity: {entity['name']}")
                current_row += 1
            
            # Fill data
            current_row = _fill_entity_rows(ws, entity, structure['column_mapping'], current_row)
            current_row += 1  # Blank row between entities
    
    # Save
    wb.save(output_path)
    return output_path


def _fill_entity_rows(ws, entity: Dict, column_mapping: Dict, start_row: int) -> int:
    """
    Fill entity field rows into worksheet.
    
    Returns the next available row index.
    """
    row_idx = start_row
    
    for field in entity.get('fields', []):
        for col_idx, field_type in column_mapping.items():
            value = _get_field_value(field, field_type)
            ws.cell(row=row_idx, column=col_idx, value=value)
        
        row_idx += 1
    
    return row_idx


def _get_field_value(field: Dict, field_type: str) -> str:
    """Extract the appropriate value from a field dict based on field type."""
    
    mapping = {
        'field_name': field.get('field_name', ''),
        'column_name': field.get('column_name', field.get('field_name', '')),
        'sql_type': field.get('sql_type', ''),
        'pk': 'PK' if field.get('pk') else '',
        'nullable': '' if field.get('nullable', True) else 'NOT NULL',
        'unique': 'UNIQUE' if field.get('unique') else '',
        'relation': field.get('relation', ''),
        'length': field.get('length', ''),
        'default': field.get('default', ''),
        'description': field.get('description', ''),
        'validation': field.get('validation', ''),
    }
    
    return str(mapping.get(field_type, ''))


def fill_word_template(entities: List[Dict], template_path: str, output_path: str) -> str:
    """
    Fill a Word template with entity metadata.
    
    Args:
        entities: List of entity dicts
        template_path: Path to template .docx file
        output_path: Path to save filled template
    
    Returns:
        Path to saved file
    """
    
    doc = Document(template_path)
    structure = TemplateMapper.detect_word_structure(doc)
    
    if structure['format'] == 'table' and structure['table_index'] is not None:
        # Fill existing table
        table = doc.tables[structure['table_index']]
        
        for entity in entities:
            # Add entity section heading if not in table
            if len(doc.paragraphs) > 0:
                doc.add_heading(entity.get('table_name', entity['name']), level=2)
            
            for field in entity.get('fields', []):
                row_cells = table.add_row().cells
                for col_idx, field_type in structure['column_mapping'].items():
                    value = _get_field_value(field, field_type)
                    row_cells[col_idx].text = value
    
    else:
        # Sections format: add entities as separate sections
        for entity in entities:
            # Add heading
            doc.add_heading(f"{entity.get('table_name', entity['name'])}", level=2)
            
            # Add a table for fields
            num_fields = len(entity.get('fields', []))
            if num_fields > 0:
                # Create table with headers
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Light Grid Accent 1'
                
                headers = ['Field Name', 'Type', 'Nullable', 'Key', 'Description']
                for idx, header in enumerate(headers):
                    table.rows[0].cells[idx].text = header
                
                # Add field rows
                for field in entity.get('fields', []):
                    row_cells = table.add_row().cells
                    row_cells[0].text = field.get('column_name', field.get('field_name', ''))
                    row_cells[1].text = field.get('sql_type', '')
                    row_cells[2].text = '' if field.get('nullable', True) else 'NOT NULL'
                    
                    key_info = []
                    if field.get('pk'):
                        key_info.append('PK')
                    if field.get('relation'):
                        key_info.append('FK')
                    row_cells[3].text = ', '.join(key_info)
                    
                    row_cells[4].text = field.get('description', '')
            
            doc.add_paragraph()  # Blank paragraph between entities
    
    doc.save(output_path)
    return output_path


def fill_template(entities: List[Dict], template_path: str, output_path: str,
                 user_column_mapping: Optional[Dict] = None) -> str:
    """
    Main entry point: auto-detect template type and fill it.
    
    Args:
        entities: List of entity dicts
        template_path: Path to template file (.xlsx or .docx)
        output_path: Path to save filled template
        user_column_mapping: Optional manual column mapping for Excel
    
    Returns:
        Path to saved file
    
    Raises:
        ValueError: If template type is not supported or structure cannot be detected
    """
    
    if template_path.endswith('.xlsx'):
        return fill_excel_template(entities, template_path, output_path, user_column_mapping)
    elif template_path.endswith('.docx'):
        return fill_word_template(entities, template_path, output_path)
    else:
        raise ValueError(f"Unsupported template format: {template_path}. Must be .xlsx or .docx")


def print_detected_structure(template_path: str) -> None:
    """
    Helper function to print the detected structure of a template.
    Useful for debugging or showing users what was detected.
    """
    
    if template_path.endswith('.xlsx'):
        wb = load_workbook(template_path, data_only=False)
        ws = wb.active
        structure = TemplateMapper.detect_excel_structure(ws)
        
        print("=== Excel Template Structure ===")
        print(f"Header Row: {structure['header_row']}")
        print(f"Data Start Row: {structure['data_start_row']}")
        print(f"Column Mapping:")
        for col_idx, field_type in sorted(structure['column_mapping'].items()):
            col_letter = ws.cell(row=1, column=col_idx).column_letter
            header_text = ws.cell(row=structure['header_row'], column=col_idx).value
            print(f"  Column {col_letter} (#{col_idx}): '{header_text}' → {field_type}")
    
    elif template_path.endswith('.docx'):
        doc = Document(template_path)
        structure = TemplateMapper.detect_word_structure(doc)
        
        print("=== Word Template Structure ===")
        print(f"Format: {structure['format']}")
        if structure['format'] == 'table':
            print(f"Table Index: {structure['table_index']}")
            print(f"Column Mapping:")
            for col_idx, field_type in sorted(structure['column_mapping'].items()):
                print(f"  Column {col_idx}: {field_type}")


if __name__ == "__main__":
    # Example usage
    sample_entities = [
        {
            "name": "User",
            "table_name": "users",
            "fields": [
                {
                    "field_name": "id",
                    "column_name": "id",
                    "sql_type": "BIGINT",
                    "pk": True,
                    "nullable": False,
                    "unique": True,
                    "description": "Primary key"
                },
                {
                    "field_name": "email",
                    "column_name": "email",
                    "sql_type": "VARCHAR(255)",
                    "pk": False,
                    "nullable": False,
                    "unique": True,
                    "length": "255",
                    "description": "User email address"
                }
            ]
        }
    ]
    
    print("Generic template filler - ready to use")
    print("Supports: .xlsx and .docx templates")
    print("Auto-detects column structure intelligently")
